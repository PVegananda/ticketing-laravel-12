"""
Script untuk membuat dokumentasi DOCX laporan tugas BK 2026
Implementasi Fitur Manajemen Event (CRUD) - Laravel 12

Nama  : Pasyah Wahyu Insanni Vegananda
NIM   : A11.2025.16575
Skema : Web Developer
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_bg(cell, hex_color):
    """Helper: set background color pada cell tabel"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    """Helper: tambah heading dengan warna opsional"""
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return h

def add_paragraph(doc, text, bold=False, italic=False, size=11, color=None):
    """Helper: tambah paragraf biasa"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def add_code_block(doc, code_text):
    """Helper: tambah blok kode dengan style monospace + background abu"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Set background style
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F3F4F6')  # abu muda
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)  # gelap
    return p

def add_screenshot_placeholder(doc, label):
    """Helper: tambah kotak placeholder screenshot"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Buat tabel 1 sel sebagai kotak
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'DBEAFE')  # biru muda
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cp.add_run(f'[ SCREENSHOT: {label} ]')
    run.bold = True
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)  # biru tua
    run.font.size = Pt(11)
    # Tambah baris kosong setelah tabel
    doc.add_paragraph()

# ============================================================
# MULAI BUAT DOKUMEN
# ============================================================
doc = Document()

# Atur margin halaman
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ============================================================
# HALAMAN COVER
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Judul utama
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run('LAPORAN TUGAS BK 2026')
title_run.bold = True
title_run.font.size = Pt(20)
title_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)

title_p2 = doc.add_paragraph()
title_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run2 = title_p2.add_run('Implementasi Fitur Manajemen Event (CRUD)')
title_run2.bold = True
title_run2.font.size = Pt(16)
title_run2.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

doc.add_paragraph()

subtitle_p = doc.add_paragraph()
subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_r = subtitle_p.add_run('Laravel 12 · Ticketing App')
subtitle_r.italic = True
subtitle_r.font.size = Pt(13)
subtitle_r.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()
doc.add_paragraph()

# Tabel identitas peserta
id_table = doc.add_table(rows=4, cols=2)
id_table.style = 'Table Grid'

data = [
    ('Nama',  'Pasyah Wahyu Insanni Vegananda'),
    ('NIM',   'A11.2025.16575'),
    ('Skema', 'Web Developer'),
    ('Tahun', '2026'),
]

for i, (label, val) in enumerate(data):
    row = id_table.rows[i]
    set_cell_bg(row.cells[0], '1E3A8A')
    set_cell_bg(row.cells[1], 'EFF6FF')

    # Label
    p0 = row.cells[0].paragraphs[0]
    r0 = p0.add_run(label)
    r0.bold = True
    r0.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r0.font.size = Pt(11)

    # Nilai
    p1 = row.cells[1].paragraphs[0]
    r1 = p1.add_run(val)
    r1.font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

halaman_baru = doc.add_paragraph()
from docx.oxml import OxmlElement
run_pb = halaman_baru.add_run()
br = OxmlElement('w:br')
br.set(qn('w:type'), 'page')
run_pb._r.append(br)

# ============================================================
# BAB 1 - HASIL IMPLEMENTASI
# ============================================================
add_heading(doc, 'BAB 1: HASIL IMPLEMENTASI', level=1, color='1E3A8A')
add_paragraph(doc,
    'Berikut adalah hasil implementasi fitur Manajemen Event (CRUD) pada aplikasi Ticketing Laravel 12. '
    'Semua perubahan dilakukan di branch manajemen-event dan di-commit satu file per commit.',
    size=11)
doc.add_paragraph()

# --- 1.1 Struktur Branch ---
add_heading(doc, '1.1 Struktur Git Branch & Commit', level=2, color='1D4ED8')
add_paragraph(doc, 'Branch yang digunakan: manajemen-event', bold=True)
add_paragraph(doc, 'Daftar commit yang dibuat (urutan terbaru):')

commits = [
    ('tambah comment di create.blade.php',        'Form tambah event baru'),
    ('tambah comment di index.blade.php',          'Daftar event + filter admin'),
    ('tambah comment di show.blade.php',           'Detail event publik + related events'),
    ('tambah comment lengkap di EventController',  'Controller CRUD event'),
    ('tambah comment penjelasan di routes/web.php','Konfigurasi route'),
    ('tambah route admin untuk manajemen event',   'Route admin events (Task 6.1)'),
    ('tambah section event terkait di show',       'Related events section (Task 5.5)'),
    ('Update edit.blade.php',                      'Form edit event + proteksi sales'),
    ('Update create.blade.php',                    'Form tambah event + dynamic tickets'),
    ('index dan admin layout',                     'Halaman index admin + sidebar menu'),
    ('Update EventController.php',                 'Logic CRUD + destroy protection'),
    ('Update Order.php',                           'Model Order untuk cek penjualan'),
    ('membuat event form request',                 'Validasi input event'),
    ('Update Event.php',                           'Model Event + accessor + scope'),
]

tbl_c = doc.add_table(rows=len(commits)+1, cols=2)
tbl_c.style = 'Table Grid'
# Header
hdr = tbl_c.rows[0]
set_cell_bg(hdr.cells[0], '1E3A8A')
set_cell_bg(hdr.cells[1], '1E3A8A')
for cell, txt in zip(hdr.cells, ['Pesan Commit', 'Keterangan']):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    r.bold = True
    r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    r.font.size = Pt(10)

for i, (msg, ket) in enumerate(commits):
    row = tbl_c.rows[i+1]
    bg = 'EFF6FF' if i % 2 == 0 else 'FFFFFF'
    set_cell_bg(row.cells[0], bg)
    set_cell_bg(row.cells[1], bg)
    row.cells[0].paragraphs[0].add_run(msg).font.size = Pt(9)
    row.cells[1].paragraphs[0].add_run(ket).font.size = Pt(9)

doc.add_paragraph()

# --- 1.2 File yang Diubah ---
add_heading(doc, '1.2 File yang Diubah / Dibuat', level=2, color='1D4ED8')

files_info = [
    ('routes/web.php',
     'Konfigurasi route publik dan admin. Tambah route resource untuk admin events dengan prefix admin dan middleware auth+verified.',
     '/admin/events'),
    ('app/Http/Controllers/EventController.php',
     'Controller utama CRUD event. Berisi method: index (list+filter), create (form), store (simpan), edit (form edit), update (perbarui), destroy (hapus+proteksi), show (detail publik).',
     '/admin/events, /events/{id}'),
    ('resources/views/pages/admin/events/index.blade.php',
     'Halaman daftar event admin. Fitur: search, filter kategori, sort tanggal, tabel dengan status badge, tombol aksi View/Edit/Delete, pagination.',
     '/admin/events'),
    ('resources/views/pages/admin/events/create.blade.php',
     'Form tambah event baru. Fitur: upload gambar + preview, tiket dinamis (tambah/hapus via JS), validasi error feedback.',
     '/admin/events/create'),
    ('resources/views/pages/admin/events/edit.blade.php',
     'Form edit event. Fitur: pre-populate data, gambar lama ditampilkan, tiket existing di-load, proteksi tanggal jika sudah ada penjualan.',
     '/admin/events/{id}/edit'),
    ('resources/views/events/show.blade.php',
     'Halaman detail event publik. Fitur: info event lengkap, daftar tiket dengan status stok, section Event Terkait (same category, upcoming, max 4).',
     '/events/{id}'),
]

for fname, desc, url in files_info:
    p = doc.add_paragraph()
    p.style = 'List Bullet'
    run_f = p.add_run(fname)
    run_f.bold = True
    run_f.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    run_d = p.add_run(f'\n   {desc}')
    run_d.font.size = Pt(10)
    run_u = p.add_run(f'\n   URL: {url}')
    run_u.italic = True
    run_u.font.size = Pt(9)
    run_u.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()

# --- 1.3 Screenshot ---
add_heading(doc, '1.3 Screenshot Hasil (Isi Sendiri)', level=2, color='1D4ED8')
add_paragraph(doc,
    'Silakan screenshot halaman berikut dari browser dan tempelkan di masing-masing kotak:',
    italic=True, color='6B7280')
doc.add_paragraph()

screenshots = [
    ('Halaman Daftar Event Admin',   'Buka: http://localhost:8000/admin/events'),
    ('Form Tambah Event',            'Buka: http://localhost:8000/admin/events/create → isi form'),
    ('Event Berhasil Dibuat',        'Setelah submit create form → tampilan success message'),
    ('Form Edit Event',              'Klik Edit pada salah satu event'),
    ('Event Berhasil Diupdate',      'Setelah submit edit form → tampilan success message'),
    ('Delete Event Biasa',           'Klik Delete → konfirmasi → success'),
    ('Delete Event (ada penjualan)', 'Coba hapus event yang punya order → tampil error message'),
    ('Halaman Detail Event Publik',  'Buka: http://localhost:8000/events/{id}'),
    ('Section Event Terkait',        'Scroll ke bawah halaman detail event → section Event Terkait'),
    ('Filter & Search Event',        'Di /admin/events → isi search, pilih kategori, sort → klik Filter'),
]

for label, hint in screenshots:
    add_paragraph(doc, f'📸 {label}', bold=True, size=11)
    add_paragraph(doc, f'   ➜ {hint}', italic=True, size=9, color='6B7280')
    add_screenshot_placeholder(doc, label)

# ============================================================
# BAB 2 - KENDALA & SOLUSI
# ============================================================
# Page break
p_br = doc.add_paragraph()
run_br2 = p_br.add_run()
br2 = OxmlElement('w:br')
br2.set(qn('w:type'), 'page')
run_br2._r.append(br2)

add_heading(doc, 'BAB 2: KENDALA ERROR & CARA MENGHINDARINYA', level=1, color='1E3A8A')
add_paragraph(doc,
    'Berikut adalah error umum yang bisa muncul saat mengerjakan tugas ini, '
    'beserta penyebab dan cara mengatasinya dalam bentuk kode.',
    size=11)
doc.add_paragraph()

# Daftar kendala
errors = [
    {
        'no': '1',
        'judul': 'Route Not Found (404) - Admin Events',
        'gejala': 'Buka /admin/events → halaman 404 Not Found',
        'penyebab': 'Route belum didaftarkan di routes/web.php atau nama route salah',
        'salah': '''// ❌ SALAH: Route tidak ada sama sekali, atau prefix/name salah
Route::get('/events', [EventController::class, 'index']);
// Nama route jadi 'events.index' bukan 'admin.events.index'
// Middleware auth tidak ada → siapa saja bisa akses''',
        'benar': '''// ✅ BENAR: Gunakan prefix 'admin', name 'admin.', dan middleware
Route::prefix('admin')
    ->name('admin.')
    ->middleware(['auth', 'verified'])
    ->group(function () {
        // Route resource otomatis buat: admin.events.index,
        // admin.events.create, admin.events.store, dll
        Route::resource('events', EventController::class)
            ->except('show'); // show ada di route publik
    });''',
    },
    {
        'no': '2',
        'judul': 'Event Gagal Dibuat - Validasi Tiket',
        'gejala': 'Submit form create event → redirect balik tanpa pesan jelas',
        'penyebab': 'Form tidak mengirim tiket sama sekali (tiket wajib minimal 1)',
        'salah': '''// ❌ SALAH: Lupa klik "Tambah Tiket" sebelum submit
// atau EventFormRequest rules tidak benar

// Di EventFormRequest.php (salah):
public function rules(): array {
    return [
        'judul' => 'required',
        // tikets tidak divalidasi → error tidak jelas
    ];
}''',
        'benar': '''// ✅ BENAR: Tambahkan validasi tikets sebagai array
public function rules(): array {
    return [
        'judul'             => 'required|string|max:255',
        'kategori_id'       => 'required|exists:kategoris,id',
        'lokasi'            => 'required|string|max:255',
        'tanggal_waktu'     => 'required|date|after:now',
        'deskripsi'         => 'required|string',
        'gambar'            => 'nullable|image|max:2048',
        'tikets'            => 'required|array|min:1',
        'tikets.*.tipe'     => 'required|string',
        'tikets.*.harga'    => 'required|numeric|min:0',
        'tikets.*.stok'     => 'required|integer|min:0',
    ];
}
// Pastikan klik "+ Tambah Tiket" minimal 1x sebelum submit form!''',
    },
    {
        'no': '3',
        'judul': 'Gambar Tidak Muncul Setelah Upload',
        'gejala': 'Event dibuat, tapi gambar tidak tampil (broken image)',
        'penyebab': 'Storage link belum dibuat atau path gambar tidak benar',
        'salah': '''// ❌ SALAH: Langsung pakai asset() tanpa cek storage link

// Di blade:
<img src="{{ asset('storage/' . $event->gambar) }}">
// Error jika php artisan storage:link belum dijalankan!

// Di controller, lupa hapus gambar lama saat update:
$validated['gambar'] = $request->file('gambar')->store('events', 'public');
// → gambar lama masih ada, storage penuh''',
        'benar': '''// ✅ LANGKAH 1: Jalankan storage link (sekali saja)
// php artisan storage:link

// ✅ LANGKAH 2: Gunakan accessor di Model Event
public function getImageUrlAttribute(): string {
    if ($this->gambar && filter_var($this->gambar, FILTER_VALIDATE_URL)) {
        return $this->gambar; // URL eksternal
    }
    $name = (!empty($this->gambar) &&
             Storage::disk('public')->exists($this->gambar))
        ? $this->gambar : 'konser.jpg';
    return Storage::url($name);
}

// ✅ LANGKAH 3: Di controller, hapus gambar lama sebelum upload baru
if ($request->hasFile('gambar')) {
    if ($event->gambar && $event->gambar !== 'konser.jpg') {
        Storage::disk('public')->delete($event->gambar);
    }
    $validated['gambar'] = $request->file('gambar')
        ->store('events', 'public');
}''',
    },
    {
        'no': '4',
        'judul': 'Event Terkait Tidak Muncul di Detail Event',
        'gejala': 'Buka /events/{id} → section "Event Terkait" kosong / error',
        'penyebab': 'Scope upcoming() belum ada di model, atau controller tidak kirim $relatedEvents',
        'salah': '''// ❌ SALAH: Controller show tidak mengirim relatedEvents ke view
public function show(Event $event) {
    $event->load(['kategori', 'tikets']);
    return view('events.show', ['event' => $event]);
    // $relatedEvents tidak ada → view error atau section tidak muncul
}

// ❌ SALAH: Scope upcoming() tidak ada di Model Event
// → error "Call to undefined method upcoming()"''',
        'benar': '''// ✅ BENAR: Model Event harus punya scope upcoming()
// app/Models/Event.php
public function scopeUpcoming($query) {
    return $query->where('tanggal_waktu', '>', now());
}

// ✅ BENAR: Controller show kirim relatedEvents
public function show(Event $event) {
    $event->load(['kategori', 'tikets']);
    $relatedEvents = Event::with('kategori')
        ->where('kategori_id', $event->kategori_id) // same category
        ->where('id', '!=', $event->id)              // bukan event ini
        ->upcoming()                                  // yang akan datang
        ->limit(4)
        ->get();
    return view('events.show', [
        'event'         => $event,
        'relatedEvents' => $relatedEvents, // wajib ada!
    ]);
}''',
    },
    {
        'no': '5',
        'judul': 'Delete Event Tidak Terlindungi (Ada Penjualan)',
        'gejala': 'Event yang sudah ada order bisa dihapus → data order jadi orphan',
        'penyebab': 'Method destroy() tidak cek hasSales() sebelum hapus',
        'salah': '''// ❌ SALAH: Langsung hapus tanpa cek penjualan
public function destroy(Event $event) {
    $event->delete(); // berbahaya! jika ada order, data jadi rusak
    return redirect()->route('admin.events.index')
        ->with('success', 'Event dihapus.');
}''',
        'benar': '''// ✅ BENAR: Cek hasSales() dulu sebelum hapus
public function destroy(Event $event) {
    // Tolak jika event punya penjualan tiket
    if ($event->hasSales()) {
        return redirect()->route('admin.events.index')
            ->with('error', 'Event tidak dapat dihapus karena sudah ada penjualan.');
    }
    // Hapus gambar dari storage (jika bukan default)
    if ($event->gambar && $event->gambar !== 'konser.jpg') {
        Storage::disk('public')->delete($event->gambar);
    }
    $event->delete();
    return redirect()->route('admin.events.index')
        ->with('success', 'Event berhasil dihapus.');
}

// Di Model Event.php, pastikan method hasSales() ada:
public function hasSales(): bool {
    return $this->orders()->exists();
    // atau: return $this->tikets()->whereHas('orders')->exists();
}''',
    },
    {
        'no': '6',
        'judul': 'CSRF Token Mismatch (419 Error)',
        'gejala': 'Submit form → error 419 Page Expired',
        'penyebab': 'Lupa @csrf di dalam form, atau session expired',
        'salah': '''{{-- ❌ SALAH: Form tanpa @csrf --}}
<form action="{{ route('admin.events.store') }}" method="POST">
    {{-- Tidak ada @csrf! --}}
    <input name="judul" type="text">
    <button type="submit">Simpan</button>
</form>

{{-- ❌ SALAH: Form delete tanpa @method('DELETE') --}}
<form action="{{ route('admin.events.destroy', $event) }}" method="POST">
    @csrf
    {{-- Lupa @method('DELETE')! --}}
    <button type="submit">Hapus</button>
</form>''',
        'benar': '''{{-- ✅ BENAR: Selalu ada @csrf di setiap form POST --}}
<form action="{{ route('admin.events.store') }}" method="POST"
      enctype="multipart/form-data">
    @csrf {{-- WAJIB: Token keamanan Laravel --}}
    ...
</form>

{{-- ✅ BENAR: Delete form pakai method spoofing --}}
<form action="{{ route('admin.events.destroy', $event) }}" method="POST">
    @csrf              {{-- Token keamanan --}}
    @method('DELETE')  {{-- Spoofing: kirim sebagai HTTP DELETE --}}
    <button class="btn btn-error">Hapus</button>
</form>''',
    },
]

for err in errors:
    # Sub-heading error
    p_no = doc.add_paragraph()
    r_no = p_no.add_run(f'Error #{err["no"]}: {err["judul"]}')
    r_no.bold = True
    r_no.font.size = Pt(12)
    r_no.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)  # merah

    # Gejala
    add_paragraph(doc, f'📋 Gejala: {err["gejala"]}', size=10, italic=True)

    # Penyebab
    add_paragraph(doc, f'⚠️  Penyebab: {err["penyebab"]}', size=10)

    # Kode salah
    add_paragraph(doc, '❌  Kode yang Salah:', bold=True, size=10, color='DC2626')
    add_code_block(doc, err['salah'])

    # Kode benar
    add_paragraph(doc, '✅  Kode yang Benar:', bold=True, size=10, color='16A34A')
    add_code_block(doc, err['benar'])

    doc.add_paragraph()

# ============================================================
# SIMPAN FILE
# ============================================================
output_path = '/Users/pvegananda/Dev/05-Workspaces/01-Core/ticketing-laravel-12/Dokumentasi_BK2026_Pasyah.docx'
doc.save(output_path)
print(f'✅ Dokumen berhasil dibuat: {output_path}')
